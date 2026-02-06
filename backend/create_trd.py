#!/usr/bin/env python3
"""
Generate Technical Requirements Document (TRD) for NEXUS Application
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph"""
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    if italic:
        p.runs[0].italic = True
    return p

def add_table_row(table, row_data):
    """Add a row to a table"""
    row = table.add_row()
    for i, cell_text in enumerate(row_data):
        row.cells[i].text = str(cell_text)

def create_trd():
    """Create the TRD document"""
    doc = Document()

    # Title Page
    title = doc.add_heading('NEXUS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Traveler Management System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)

    doc.add_paragraph()

    trd_title = doc.add_paragraph('Technical Requirements Document (TRD)')
    trd_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trd_title.runs[0].bold = True
    trd_title.runs[0].font.size = Pt(16)

    doc.add_paragraph()
    doc.add_paragraph()

    company = doc.add_paragraph('American Circuits')
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company.runs[0].font.size = Pt(14)

    date = doc.add_paragraph(f'Document Date: {datetime.now().strftime("%B %d, %Y")}')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER

    version = doc.add_paragraph('Version: 1.0')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    add_heading(doc, 'Table of Contents', level=1)
    doc.add_paragraph('1. Executive Summary')
    doc.add_paragraph('2. System Overview')
    doc.add_paragraph('3. Architecture')
    doc.add_paragraph('4. Technical Stack')
    doc.add_paragraph('5. Features & Functionality')
    doc.add_paragraph('6. User Roles & Permissions')
    doc.add_paragraph('7. API Endpoints')
    doc.add_paragraph('8. Database Schema')
    doc.add_paragraph('9. Security & Authentication')
    doc.add_paragraph('10. Deployment & Infrastructure')
    doc.add_paragraph('11. Future Enhancements')
    doc.add_paragraph('12. Appendix')

    doc.add_page_break()

    # 1. Executive Summary
    add_heading(doc, '1. Executive Summary', level=1)
    add_paragraph(doc,
        'NEXUS is a comprehensive Traveler Management System designed for American Circuits to streamline '
        'manufacturing operations, track work orders, manage labor hours, and maintain quality control throughout '
        'the production lifecycle. The system provides real-time visibility into production status, barcode/QR code '
        'tracking, and automated workflow management.')

    add_heading(doc, '1.1 Purpose', level=2)
    add_paragraph(doc,
        'The purpose of this document is to provide a comprehensive technical specification of the NEXUS application, '
        'including its architecture, features, technical requirements, and implementation details.')

    add_heading(doc, '1.2 Scope', level=2)
    add_paragraph(doc,
        'This document covers the complete NEXUS system including:')
    doc.add_paragraph('• Frontend web application (Next.js)', style='List Bullet')
    doc.add_paragraph('• Backend API server (FastAPI)', style='List Bullet')
    doc.add_paragraph('• Database layer (PostgreSQL)', style='List Bullet')
    doc.add_paragraph('• Authentication and authorization', style='List Bullet')
    doc.add_paragraph('• Barcode/QR code generation and scanning', style='List Bullet')
    doc.add_paragraph('• Reporting and analytics', style='List Bullet')

    doc.add_page_break()

    # 2. System Overview
    add_heading(doc, '2. System Overview', level=1)

    add_heading(doc, '2.1 System Purpose', level=2)
    add_paragraph(doc,
        'NEXUS is designed to digitize and automate the traveler (work order) process in manufacturing. '
        'It replaces manual paper-based travelers with a digital system that provides:')
    doc.add_paragraph('• Real-time tracking of work orders through production stages', style='List Bullet')
    doc.add_paragraph('• Automated barcode and QR code generation for easy identification', style='List Bullet')
    doc.add_paragraph('• Labor hour tracking and reporting', style='List Bullet')
    doc.add_paragraph('• Quality control with acceptance/rejection tracking', style='List Bullet')
    doc.add_paragraph('• Digital signatures and audit trails', style='List Bullet')
    doc.add_paragraph('• Multi-step process routing with work center management', style='List Bullet')

    add_heading(doc, '2.2 Key Features', level=2)
    doc.add_paragraph('• Dashboard with production metrics and KPIs', style='List Bullet')
    doc.add_paragraph('• Traveler creation with multiple types (PCB, ASSY, CABLE, etc.)', style='List Bullet')
    doc.add_paragraph('• Work center routing with customizable process steps', style='List Bullet')
    doc.add_paragraph('• Barcode/QR code scanning for tracking', style='List Bullet')
    doc.add_paragraph('• Labor tracking with clock in/out functionality', style='List Bullet')
    doc.add_paragraph('• Approval workflow for supervisors/managers', style='List Bullet')
    doc.add_paragraph('• Comprehensive reporting and analytics', style='List Bullet')
    doc.add_paragraph('• Active/Inactive traveler management', style='List Bullet')
    doc.add_paragraph('• Print-optimized traveler documents', style='List Bullet')

    doc.add_page_break()

    # 3. Architecture
    add_heading(doc, '3. Architecture', level=1)

    add_heading(doc, '3.1 System Architecture', level=2)
    add_paragraph(doc,
        'NEXUS follows a modern three-tier architecture:')

    doc.add_paragraph('• Presentation Layer: Next.js 15 with React and TypeScript', style='List Bullet')
    doc.add_paragraph('• Business Logic Layer: FastAPI (Python) REST API', style='List Bullet')
    doc.add_paragraph('• Data Layer: PostgreSQL relational database', style='List Bullet')
    doc.add_paragraph('• Reverse Proxy: NGINX for request routing and load balancing', style='List Bullet')

    add_heading(doc, '3.2 Component Diagram', level=2)
    add_paragraph(doc, 'The system consists of the following components:')

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Component'
    header_cells[1].text = 'Technology'
    header_cells[2].text = 'Purpose'

    components = [
        ('Frontend', 'Next.js 15 + TypeScript', 'User interface and client-side logic'),
        ('Backend API', 'FastAPI + Python 3.11', 'Business logic and data processing'),
        ('Database', 'PostgreSQL 15', 'Persistent data storage'),
        ('Reverse Proxy', 'NGINX', 'Request routing and static file serving'),
        ('Containerization', 'Docker + Docker Compose', 'Application deployment and orchestration'),
    ]

    for comp in components:
        add_table_row(table, comp)

    add_heading(doc, '3.3 Deployment Architecture', level=2)
    add_paragraph(doc,
        'NEXUS is containerized using Docker and orchestrated with Docker Compose. The deployment consists of:')
    doc.add_paragraph('• nexus_frontend: Next.js application (Port 3000)', style='List Bullet')
    doc.add_paragraph('• nexus_backend: FastAPI server (Port 8000)', style='List Bullet')
    doc.add_paragraph('• nexus_postgres: PostgreSQL database (Port 5432)', style='List Bullet')
    doc.add_paragraph('• nexus_nginx: NGINX reverse proxy (Port 100)', style='List Bullet')

    doc.add_page_break()

    # 4. Technical Stack
    add_heading(doc, '4. Technical Stack', level=1)

    add_heading(doc, '4.1 Frontend Technologies', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Technology'
    header_cells[1].text = 'Version'
    header_cells[2].text = 'Purpose'

    frontend_tech = [
        ('Next.js', '15.5.4', 'React framework with SSR/SSG'),
        ('React', '19.x', 'UI component library'),
        ('TypeScript', '5.x', 'Type-safe JavaScript'),
        ('Tailwind CSS', '3.x', 'Utility-first CSS framework'),
        ('Heroicons', '2.x', 'Icon library'),
    ]

    for tech in frontend_tech:
        add_table_row(table, tech)

    add_heading(doc, '4.2 Backend Technologies', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Technology'
    header_cells[1].text = 'Version'
    header_cells[2].text = 'Purpose'

    backend_tech = [
        ('FastAPI', '0.115.x', 'Modern Python web framework'),
        ('Python', '3.11', 'Programming language'),
        ('SQLAlchemy', '2.x', 'ORM for database operations'),
        ('Alembic', 'Latest', 'Database migration tool'),
        ('Pydantic', '2.x', 'Data validation'),
        ('Passlib', 'Latest', 'Password hashing'),
        ('python-jose', 'Latest', 'JWT token handling'),
        ('python-multipart', 'Latest', 'File upload handling'),
    ]

    for tech in backend_tech:
        add_table_row(table, tech)

    add_heading(doc, '4.3 Database', level=2)
    add_paragraph(doc, 'PostgreSQL 15 - Enterprise-grade relational database')

    add_heading(doc, '4.4 DevOps & Infrastructure', level=2)
    doc.add_paragraph('• Docker: Containerization', style='List Bullet')
    doc.add_paragraph('• Docker Compose: Multi-container orchestration', style='List Bullet')
    doc.add_paragraph('• NGINX: Reverse proxy and web server', style='List Bullet')

    doc.add_page_break()

    # 5. Features & Functionality
    add_heading(doc, '5. Features & Functionality', level=1)

    add_heading(doc, '5.1 Dashboard', level=2)
    add_paragraph(doc, 'The dashboard provides an overview of manufacturing operations:')
    doc.add_paragraph('• Active travelers count', style='List Bullet')
    doc.add_paragraph('• Completed travelers count', style='List Bullet')
    doc.add_paragraph('• Total parts in production', style='List Bullet')
    doc.add_paragraph('• Recent travelers list (5 most recent)', style='List Bullet')
    doc.add_paragraph('• Quick action buttons (Create Traveler, Clock In/Out, Reports)', style='List Bullet')
    doc.add_paragraph('• Simple, color-neutral design for professional appearance', style='List Bullet')

    add_heading(doc, '5.2 Traveler Management', level=2)

    add_heading(doc, '5.2.1 Traveler Creation', level=3)
    add_paragraph(doc, 'Users can create new travelers with the following information:')
    doc.add_paragraph('• Job Number (unique identifier)', style='List Bullet')
    doc.add_paragraph('• Work Order Number', style='List Bullet')
    doc.add_paragraph('• PO Number (optional)', style='List Bullet')
    doc.add_paragraph('• Traveler Type (PCB, ASSY, CABLE, PURCHASING)', style='List Bullet')
    doc.add_paragraph('• Part Number and Description', style='List Bullet')
    doc.add_paragraph('• Revision (Part, Customer)', style='List Bullet')
    doc.add_paragraph('• Quantity', style='List Bullet')
    doc.add_paragraph('• Customer Code and Name', style='List Bullet')
    doc.add_paragraph('• Due Date and Ship Date', style='List Bullet')
    doc.add_paragraph('• Specifications with dates', style='List Bullet')
    doc.add_paragraph('• Shipping information (From Stock, To Stock, Ship Via)', style='List Bullet')
    doc.add_paragraph('• Comments and Notes', style='List Bullet')
    doc.add_paragraph('• Work Center routing steps', style='List Bullet')

    add_heading(doc, '5.2.2 Work Center Routing', level=3)
    add_paragraph(doc, 'Each traveler includes a routing table with process steps:')
    doc.add_paragraph('• Sequence Number', style='List Bullet')
    doc.add_paragraph('• Work Center (predefined or custom)', style='List Bullet')
    doc.add_paragraph('• Instructions', style='List Bullet')
    doc.add_paragraph('• Time tracking', style='List Bullet')
    doc.add_paragraph('• Quantity, Accepted, Rejected counts', style='List Bullet')
    doc.add_paragraph('• Digital signature field', style='List Bullet')
    doc.add_paragraph('• Completion date', style='List Bullet')
    doc.add_paragraph('• QR code for each work center for scanning', style='List Bullet')

    add_heading(doc, '5.2.3 Predefined Work Centers', level=3)
    add_paragraph(doc, 'The system includes predefined work centers:')
    work_centers = [
        'INCOMING INSPECTION', 'PCB PREP', 'ASSEMBLY', 'SOLDERING', 'TESTING',
        'QC INSPECTION', 'PACKAGING', 'SHIPPING', 'REWORK', 'CUSTOM (user-defined)'
    ]
    for wc in work_centers:
        doc.add_paragraph(f'• {wc}', style='List Bullet')

    add_heading(doc, '5.2.4 Barcode & QR Code Generation', level=3)
    add_paragraph(doc, 'Each traveler automatically generates:')
    doc.add_paragraph('• Barcode with job number (for header scanning)', style='List Bullet')
    doc.add_paragraph('• QR code with job number (for header scanning)', style='List Bullet')
    doc.add_paragraph('• QR code for each work center step (for step-level tracking)', style='List Bullet')

    add_heading(doc, '5.2.5 Traveler Status Management', level=3)
    add_paragraph(doc, 'Travelers can have the following statuses:')
    doc.add_paragraph('• CREATED: Initial state', style='List Bullet')
    doc.add_paragraph('• IN_PROGRESS: Work has started', style='List Bullet')
    doc.add_paragraph('• COMPLETED: All steps finished', style='List Bullet')
    doc.add_paragraph('• ON_HOLD: Temporarily paused', style='List Bullet')
    doc.add_paragraph('• CANCELLED: Cancelled/voided', style='List Bullet')

    add_paragraph(doc, 'Additionally, travelers can be marked as Active or Inactive:')
    doc.add_paragraph('• Active: Visible in main lists and dashboards', style='List Bullet')
    doc.add_paragraph('• Inactive: Archived but not deleted, hidden from main views', style='List Bullet')
    doc.add_paragraph('• Any user can toggle active/inactive status', style='List Bullet')

    add_heading(doc, '5.2.6 Print Functionality', level=3)
    add_paragraph(doc, 'Travelers can be printed with optimized formatting:')
    doc.add_paragraph('• Black and white print styling', style='List Bullet')
    doc.add_paragraph('• Optimized font sizes (14px for readability)', style='List Bullet')
    doc.add_paragraph('• Proper page breaks for labor hours section', style='List Bullet')
    doc.add_paragraph('• Barcodes and QR codes included', style='List Bullet')
    doc.add_paragraph('• All routing steps and specifications', style='List Bullet')

    add_heading(doc, '5.3 Labor Tracking', level=2)
    add_paragraph(doc, 'The labor tracking module allows:')
    doc.add_paragraph('• Clock in/out functionality for operators', style='List Bullet')
    doc.add_paragraph('• Time tracking per traveler', style='List Bullet')
    doc.add_paragraph('• Labor hours reporting', style='List Bullet')
    doc.add_paragraph('• Operator name tracking', style='List Bullet')
    doc.add_paragraph('• Total hours calculation', style='List Bullet')
    doc.add_paragraph('• Optional labor hours section on travelers', style='List Bullet')

    add_heading(doc, '5.4 Barcode/QR Code Scanning', level=2)
    add_paragraph(doc, 'The scanning functionality enables:')
    doc.add_paragraph('• Scan traveler header barcode/QR code to view details', style='List Bullet')
    doc.add_paragraph('• Scan work center QR codes to track progress', style='List Bullet')
    doc.add_paragraph('• Real-time location tracking of travelers', style='List Bullet')
    doc.add_paragraph('• Scan history and audit trail', style='List Bullet')
    doc.add_paragraph('• Support for handheld barcode scanners', style='List Bullet')

    add_heading(doc, '5.5 Approval Workflow', level=2)
    add_paragraph(doc, 'Certain operations require approval:')
    doc.add_paragraph('• Supervisor/Manager approval for specific actions', style='List Bullet')
    doc.add_paragraph('• Approval queue for pending requests', style='List Bullet')
    doc.add_paragraph('• Email notifications to approvers', style='List Bullet')
    doc.add_paragraph('• Approval history and audit trail', style='List Bullet')

    add_heading(doc, '5.6 Reporting & Analytics', level=2)
    add_paragraph(doc, 'The system provides various reports:')
    doc.add_paragraph('• Production summary reports', style='List Bullet')
    doc.add_paragraph('• Labor hours reports', style='List Bullet')
    doc.add_paragraph('• Work center utilization', style='List Bullet')
    doc.add_paragraph('• Quality metrics (acceptance/rejection rates)', style='List Bullet')
    doc.add_paragraph('• Traveler history and status', style='List Bullet')

    doc.add_page_break()

    # 6. User Roles & Permissions
    add_heading(doc, '6. User Roles & Permissions', level=1)

    add_paragraph(doc, 'NEXUS implements role-based access control (RBAC) with four user roles:')

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Role'
    header_cells[1].text = 'Permissions'
    header_cells[2].text = 'Description'

    roles = [
        ('ADMIN',
         'Full system access, user management, delete travelers, approve/reject',
         'System administrators with complete control. Current admins: Adam, Kris, Alex, Preet'),
        ('SUPERVISOR',
         'Create/edit travelers, approve requests, view all data',
         'Production supervisors and managers'),
        ('OPERATOR',
         'Create travelers, update work steps, clock in/out, view assigned work',
         'Floor operators and technicians'),
        ('VIEWER',
         'Read-only access to travelers and reports',
         'View-only access for stakeholders'),
    ]

    for role in roles:
        add_table_row(table, role)

    add_heading(doc, '6.1 Admin Users', level=2)
    add_paragraph(doc, 'The following users have administrator privileges:')
    doc.add_paragraph('• Adam (username: adam)', style='List Bullet')
    doc.add_paragraph('• Kris (username: kris)', style='List Bullet')
    doc.add_paragraph('• Alex (username: alex)', style='List Bullet')
    doc.add_paragraph('• Preet (username: preet)', style='List Bullet')

    add_paragraph(doc, 'All admin users can:')
    doc.add_paragraph('• Delete travelers', style='List Bullet')
    doc.add_paragraph('• Create and manage user accounts', style='List Bullet')
    doc.add_paragraph('• Approve/reject all requests', style='List Bullet')
    doc.add_paragraph('• Access all system functionality', style='List Bullet')

    doc.add_page_break()

    # 7. API Endpoints
    add_heading(doc, '7. API Endpoints', level=1)

    add_heading(doc, '7.1 Authentication Endpoints', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Endpoint'
    header_cells[1].text = 'Method'
    header_cells[2].text = 'Description'

    auth_endpoints = [
        ('/auth/login', 'POST', 'User login, returns JWT token'),
        ('/auth/refresh', 'POST', 'Refresh JWT token'),
        ('/auth/me', 'GET', 'Get current user information'),
    ]

    for ep in auth_endpoints:
        add_table_row(table, ep)

    add_heading(doc, '7.2 Traveler Endpoints', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Endpoint'
    header_cells[1].text = 'Method'
    header_cells[2].text = 'Description'

    traveler_endpoints = [
        ('/travelers/', 'GET', 'List all travelers'),
        ('/travelers/', 'POST', 'Create new traveler'),
        ('/travelers/{id}', 'GET', 'Get traveler by ID'),
        ('/travelers/by-job/{job_number}', 'GET', 'Get traveler by job number'),
        ('/travelers/{id}', 'PUT', 'Update traveler'),
        ('/travelers/{id}', 'PATCH', 'Partial update (e.g., toggle active)'),
        ('/travelers/{id}', 'DELETE', 'Delete traveler (admin only)'),
        ('/travelers/manufacturing-steps/{type}', 'GET', 'Get predefined steps by type'),
    ]

    for ep in traveler_endpoints:
        add_table_row(table, ep)

    add_heading(doc, '7.3 Tracking Endpoints', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Endpoint'
    header_cells[1].text = 'Method'
    header_cells[2].text = 'Description'

    tracking_endpoints = [
        ('/travelers/scan', 'POST', 'Record barcode/QR scan'),
        ('/travelers/tracking/logs', 'GET', 'Get tracking logs'),
        ('/travelers/tracking/current-location/{job}', 'GET', 'Get current location'),
    ]

    for ep in tracking_endpoints:
        add_table_row(table, ep)

    add_heading(doc, '7.4 User Management Endpoints', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Endpoint'
    header_cells[1].text = 'Method'
    header_cells[2].text = 'Description'

    user_endpoints = [
        ('/users/', 'GET', 'List users (admin/supervisor)'),
        ('/users/', 'POST', 'Create user (admin only)'),
        ('/users/{id}', 'GET', 'Get user by ID'),
        ('/users/{id}', 'PUT', 'Update user'),
        ('/users/{id}', 'DELETE', 'Deactivate user (admin only)'),
        ('/users/approvers/list', 'GET', 'Get list of approvers'),
    ]

    for ep in user_endpoints:
        add_table_row(table, ep)

    add_heading(doc, '7.5 Other Endpoints', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Endpoint'
    header_cells[1].text = 'Method'
    header_cells[2].text = 'Description'

    other_endpoints = [
        ('/labor/', 'POST', 'Clock in/out for labor tracking'),
        ('/labor/', 'GET', 'Get labor entries'),
        ('/approvals/', 'GET', 'Get pending approvals'),
        ('/approvals/{id}/approve', 'POST', 'Approve request'),
        ('/approvals/{id}/reject', 'POST', 'Reject request'),
        ('/barcodes/generate', 'POST', 'Generate barcode image'),
    ]

    for ep in other_endpoints:
        add_table_row(table, ep)

    doc.add_page_break()

    # 8. Database Schema
    add_heading(doc, '8. Database Schema', level=1)

    add_heading(doc, '8.1 Main Tables', level=2)

    add_heading(doc, '8.1.1 users', level=3)
    add_paragraph(doc, 'Stores user account information')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Column'
    header_cells[1].text = 'Type'
    header_cells[2].text = 'Description'

    user_columns = [
        ('id', 'INTEGER', 'Primary key'),
        ('username', 'VARCHAR(50)', 'Unique username'),
        ('email', 'VARCHAR(100)', 'Email address'),
        ('first_name', 'VARCHAR(50)', 'First name'),
        ('last_name', 'VARCHAR(50)', 'Last name'),
        ('hashed_password', 'VARCHAR(255)', 'Bcrypt hashed password'),
        ('role', 'ENUM', 'ADMIN, SUPERVISOR, OPERATOR, VIEWER'),
        ('is_approver', 'BOOLEAN', 'Can approve requests'),
        ('is_active', 'BOOLEAN', 'Account active status'),
        ('created_at', 'TIMESTAMP', 'Account creation timestamp'),
    ]

    for col in user_columns:
        add_table_row(table, col)

    add_heading(doc, '8.1.2 travelers', level=3)
    add_paragraph(doc, 'Stores traveler/work order information')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Column'
    header_cells[1].text = 'Type'
    header_cells[2].text = 'Description'

    traveler_columns = [
        ('id', 'INTEGER', 'Primary key'),
        ('job_number', 'VARCHAR(50)', 'Unique job number'),
        ('work_order_number', 'VARCHAR(50)', 'Work order reference'),
        ('po_number', 'VARCHAR(50)', 'Purchase order number'),
        ('traveler_type', 'ENUM', 'PCB, ASSY, CABLE, etc.'),
        ('part_number', 'VARCHAR(100)', 'Part number'),
        ('part_description', 'TEXT', 'Part description'),
        ('revision', 'VARCHAR(20)', 'Revision number'),
        ('customer_revision', 'VARCHAR(20)', 'Customer revision'),
        ('part_revision', 'VARCHAR(20)', 'Part revision'),
        ('quantity', 'INTEGER', 'Order quantity'),
        ('customer_code', 'VARCHAR(50)', 'Customer code'),
        ('customer_name', 'VARCHAR(200)', 'Customer name'),
        ('status', 'ENUM', 'CREATED, IN_PROGRESS, etc.'),
        ('priority', 'ENUM', 'LOW, NORMAL, HIGH, URGENT'),
        ('work_center', 'VARCHAR(100)', 'Current work center'),
        ('specs', 'TEXT', 'Specifications (JSON)'),
        ('specs_date', 'DATE', 'Specification date'),
        ('from_stock', 'VARCHAR(100)', 'From stock location'),
        ('to_stock', 'VARCHAR(100)', 'To stock location'),
        ('ship_via', 'VARCHAR(100)', 'Shipping method'),
        ('comments', 'TEXT', 'Comments and notes'),
        ('due_date', 'DATE', 'Due date'),
        ('ship_date', 'DATE', 'Ship date'),
        ('is_active', 'BOOLEAN', 'Active/Inactive status'),
        ('include_labor_hours', 'BOOLEAN', 'Include labor tracking'),
        ('created_by', 'INTEGER', 'Foreign key to users'),
        ('created_at', 'TIMESTAMP', 'Creation timestamp'),
        ('updated_at', 'TIMESTAMP', 'Last update timestamp'),
    ]

    for col in traveler_columns:
        add_table_row(table, col)

    add_heading(doc, '8.1.3 process_steps', level=3)
    add_paragraph(doc, 'Stores routing steps for travelers')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Column'
    header_cells[1].text = 'Type'
    header_cells[2].text = 'Description'

    step_columns = [
        ('id', 'INTEGER', 'Primary key'),
        ('traveler_id', 'INTEGER', 'Foreign key to travelers'),
        ('step_number', 'INTEGER', 'Sequence number'),
        ('operation', 'VARCHAR(100)', 'Operation/work center name'),
        ('work_center_code', 'VARCHAR(50)', 'Work center code'),
        ('instructions', 'TEXT', 'Step instructions'),
        ('estimated_time', 'INTEGER', 'Estimated time (minutes)'),
        ('quantity', 'INTEGER', 'Quantity processed'),
        ('accepted', 'INTEGER', 'Accepted quantity'),
        ('rejected', 'INTEGER', 'Rejected quantity'),
        ('sign', 'VARCHAR(50)', 'Operator signature'),
        ('completed_date', 'DATE', 'Completion date'),
        ('is_completed', 'BOOLEAN', 'Completion status'),
        ('is_required', 'BOOLEAN', 'Required step'),
    ]

    for col in step_columns:
        add_table_row(table, col)

    add_heading(doc, '8.1.4 traveler_tracking_log', level=3)
    add_paragraph(doc, 'Stores barcode/QR code scan history')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Column'
    header_cells[1].text = 'Type'
    header_cells[2].text = 'Description'

    log_columns = [
        ('id', 'INTEGER', 'Primary key'),
        ('traveler_id', 'INTEGER', 'Foreign key to travelers'),
        ('job_number', 'VARCHAR(50)', 'Job number'),
        ('work_center', 'VARCHAR(100)', 'Work center scanned'),
        ('step_sequence', 'INTEGER', 'Step sequence number'),
        ('scan_type', 'VARCHAR(20)', 'HEADER or WORK_CENTER'),
        ('scanned_by', 'VARCHAR(100)', 'Operator name'),
        ('scanned_at', 'TIMESTAMP', 'Scan timestamp'),
        ('notes', 'TEXT', 'Additional notes'),
    ]

    for col in log_columns:
        add_table_row(table, col)

    add_heading(doc, '8.2 Additional Tables', level=2)
    add_paragraph(doc, 'The system also includes these supporting tables:')
    doc.add_paragraph('• work_centers: Defines available work centers', style='List Bullet')
    doc.add_paragraph('• work_orders: Work order master data', style='List Bullet')
    doc.add_paragraph('• labor_entries: Labor time tracking', style='List Bullet')
    doc.add_paragraph('• approvals: Approval workflow records', style='List Bullet')
    doc.add_paragraph('• audit_logs: System audit trail', style='List Bullet')
    doc.add_paragraph('• sub_steps: Detailed sub-steps for process steps', style='List Bullet')
    doc.add_paragraph('• manual_steps: User-added manual steps', style='List Bullet')

    doc.add_page_break()

    # 9. Security & Authentication
    add_heading(doc, '9. Security & Authentication', level=1)

    add_heading(doc, '9.1 Authentication', level=2)
    add_paragraph(doc, 'NEXUS uses JWT (JSON Web Token) based authentication:')
    doc.add_paragraph('• Users login with username/password', style='List Bullet')
    doc.add_paragraph('• Server validates credentials and issues JWT token', style='List Bullet')
    doc.add_paragraph('• Token stored in browser localStorage', style='List Bullet')
    doc.add_paragraph('• Token included in Authorization header for API requests', style='List Bullet')
    doc.add_paragraph('• Token expiration and refresh mechanism', style='List Bullet')

    add_heading(doc, '9.2 Password Security', level=2)
    add_paragraph(doc, 'Password security measures:')
    doc.add_paragraph('• Passwords hashed using bcrypt algorithm', style='List Bullet')
    doc.add_paragraph('• Salt rounds configured for security', style='List Bullet')
    doc.add_paragraph('• Plain text passwords never stored', style='List Bullet')
    doc.add_paragraph('• Password reset functionality available', style='List Bullet')

    add_heading(doc, '9.3 Authorization', level=2)
    add_paragraph(doc, 'Role-based access control (RBAC):')
    doc.add_paragraph('• Each endpoint checks user role', style='List Bullet')
    doc.add_paragraph('• Unauthorized access returns 403 Forbidden', style='List Bullet')
    doc.add_paragraph('• Admin-only operations protected', style='List Bullet')
    doc.add_paragraph('• Approval workflow enforces permissions', style='List Bullet')

    add_heading(doc, '9.4 Data Security', level=2)
    add_paragraph(doc, 'Data protection measures:')
    doc.add_paragraph('• CORS configured for allowed origins', style='List Bullet')
    doc.add_paragraph('• SQL injection protection via ORM', style='List Bullet')
    doc.add_paragraph('• Input validation using Pydantic', style='List Bullet')
    doc.add_paragraph('• Audit logging for sensitive operations', style='List Bullet')

    doc.add_page_break()

    # 10. Deployment & Infrastructure
    add_heading(doc, '10. Deployment & Infrastructure', level=1)

    add_heading(doc, '10.1 Docker Configuration', level=2)
    add_paragraph(doc, 'The application is fully containerized:')

    add_heading(doc, '10.1.1 Frontend Container', level=3)
    doc.add_paragraph('• Base Image: node:18-alpine', style='List Bullet')
    doc.add_paragraph('• Exposed Port: 3000', style='List Bullet')
    doc.add_paragraph('• Build: Next.js production build', style='List Bullet')
    doc.add_paragraph('• Environment: Production optimized', style='List Bullet')

    add_heading(doc, '10.1.2 Backend Container', level=3)
    doc.add_paragraph('• Base Image: python:3.11-slim', style='List Bullet')
    doc.add_paragraph('• Exposed Port: 8000', style='List Bullet')
    doc.add_paragraph('• Dependencies: gcc, postgresql-client', style='List Bullet')
    doc.add_paragraph('• Application Server: Uvicorn', style='List Bullet')

    add_heading(doc, '10.1.3 Database Container', level=3)
    doc.add_paragraph('• Base Image: postgres:15', style='List Bullet')
    doc.add_paragraph('• Exposed Port: 5432', style='List Bullet')
    doc.add_paragraph('• Volume: Persistent data storage', style='List Bullet')
    doc.add_paragraph('• Auto-creates database on first run', style='List Bullet')

    add_heading(doc, '10.1.4 NGINX Container', level=3)
    doc.add_paragraph('• Base Image: nginx:alpine', style='List Bullet')
    doc.add_paragraph('• Exposed Port: 100 (external)', style='List Bullet')
    doc.add_paragraph('• Routes API requests to backend', style='List Bullet')
    doc.add_paragraph('• Serves frontend static files', style='List Bullet')

    add_heading(doc, '10.2 Network Configuration', level=2)
    add_paragraph(doc, 'Docker network setup:')
    doc.add_paragraph('• Custom bridge network: nexus_network', style='List Bullet')
    doc.add_paragraph('• Internal DNS resolution between containers', style='List Bullet')
    doc.add_paragraph('• Only NGINX exposed to external network', style='List Bullet')

    add_heading(doc, '10.3 Data Persistence', level=2)
    add_paragraph(doc, 'Data is persisted through:')
    doc.add_paragraph('• PostgreSQL volume: postgres_data', style='List Bullet')
    doc.add_paragraph('• File uploads: backend static directory', style='List Bullet')
    doc.add_paragraph('• Database backups: manual or automated', style='List Bullet')

    add_heading(doc, '10.4 Startup & Seeding', level=2)
    add_paragraph(doc, 'On first startup:')
    doc.add_paragraph('• Database tables automatically created', style='List Bullet')
    doc.add_paragraph('• Admin users seeded (Adam, Kris, Alex, Preet)', style='List Bullet')
    doc.add_paragraph('• Default password: admin123', style='List Bullet')
    doc.add_paragraph('• Sample travelers optionally seeded', style='List Bullet')

    add_heading(doc, '10.5 Deployment Commands', level=2)
    add_paragraph(doc, 'Common deployment commands:')
    doc.add_paragraph('• Build and start: docker-compose up -d --build', style='List Bullet')
    doc.add_paragraph('• Stop all: docker-compose down', style='List Bullet')
    doc.add_paragraph('• View logs: docker-compose logs -f [service]', style='List Bullet')
    doc.add_paragraph('• Restart service: docker-compose restart [service]', style='List Bullet')

    doc.add_page_break()

    # 11. Future Enhancements
    add_heading(doc, '11. Future Enhancements', level=1)

    add_paragraph(doc, 'Potential future enhancements for NEXUS:')

    add_heading(doc, '11.1 Advanced Reporting', level=2)
    doc.add_paragraph('• Interactive dashboards with charts and graphs', style='List Bullet')
    doc.add_paragraph('• Export reports to PDF, Excel, CSV', style='List Bullet')
    doc.add_paragraph('• Custom report builder', style='List Bullet')
    doc.add_paragraph('• Scheduled report generation and email delivery', style='List Bullet')

    add_heading(doc, '11.2 Mobile Application', level=2)
    doc.add_paragraph('• Native iOS/Android apps for operators', style='List Bullet')
    doc.add_paragraph('• Mobile barcode scanning with camera', style='List Bullet')
    doc.add_paragraph('• Offline mode with sync', style='List Bullet')
    doc.add_paragraph('• Push notifications for updates', style='List Bullet')

    add_heading(doc, '11.3 Integration', level=2)
    doc.add_paragraph('• ERP system integration (SAP, Oracle, etc.)', style='List Bullet')
    doc.add_paragraph('• Email system integration for automated notifications', style='List Bullet')
    doc.add_paragraph('• External barcode label printer integration', style='List Bullet')
    doc.add_paragraph('• Inventory management system connection', style='List Bullet')

    add_heading(doc, '11.4 Advanced Features', level=2)
    doc.add_paragraph('• Real-time notifications and alerts', style='List Bullet')
    doc.add_paragraph('• Digital signatures with e-signature pad support', style='List Bullet')
    doc.add_paragraph('• Photo attachments for quality documentation', style='List Bullet')
    doc.add_paragraph('• Video instructions for complex operations', style='List Bullet')
    doc.add_paragraph('• Machine integration for automated data collection', style='List Bullet')
    doc.add_paragraph('• AI-powered quality inspection', style='List Bullet')

    add_heading(doc, '11.5 Performance Optimization', level=2)
    doc.add_paragraph('• Caching layer (Redis)', style='List Bullet')
    doc.add_paragraph('• Database query optimization', style='List Bullet')
    doc.add_paragraph('• CDN for static assets', style='List Bullet')
    doc.add_paragraph('• Load balancing for high traffic', style='List Bullet')

    doc.add_page_break()

    # 12. Appendix
    add_heading(doc, '12. Appendix', level=1)

    add_heading(doc, '12.1 Glossary', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Term'
    header_cells[1].text = 'Definition'

    glossary = [
        ('Traveler', 'A document that travels with a work order through production'),
        ('Work Center', 'A physical location or station in the manufacturing process'),
        ('Routing', 'The sequence of operations required to complete a product'),
        ('QR Code', 'Quick Response code - 2D barcode for rapid scanning'),
        ('JWT', 'JSON Web Token - authentication token standard'),
        ('ORM', 'Object-Relational Mapping - database abstraction layer'),
        ('RBAC', 'Role-Based Access Control - permission system'),
        ('SSR', 'Server-Side Rendering - rendering web pages on server'),
        ('API', 'Application Programming Interface'),
        ('PCB', 'Printed Circuit Board'),
    ]

    for term in glossary:
        add_table_row(table, term)

    add_heading(doc, '12.2 System URLs', level=2)
    add_paragraph(doc, 'Access the system at:')
    doc.add_paragraph('• Production: http://acidashboard.aci.local:100', style='List Bullet')
    doc.add_paragraph('• API Documentation: http://acidashboard.aci.local:100/api/docs', style='List Bullet')
    doc.add_paragraph('• API Backend: http://acidashboard.aci.local:100/api', style='List Bullet')

    add_heading(doc, '12.3 Contact Information', level=2)
    add_paragraph(doc, 'For support or questions about NEXUS:')
    doc.add_paragraph('• System Administrators: Adam, Kris, Alex, Preet', style='List Bullet')
    doc.add_paragraph('• Company: American Circuits', style='List Bullet')

    add_heading(doc, '12.4 Document Revision History', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Version'
    header_cells[1].text = 'Date'
    header_cells[2].text = 'Author'
    header_cells[3].text = 'Changes'

    revisions = [
        ('1.0', datetime.now().strftime("%Y-%m-%d"), 'Claude AI', 'Initial TRD creation'),
    ]

    for rev in revisions:
        add_table_row(table, rev)

    # Save document
    doc.save('/app/NEXUS_TRD.docx')
    print("✅ TRD document created successfully: NEXUS_TRD.docx")

if __name__ == '__main__':
    create_trd()
